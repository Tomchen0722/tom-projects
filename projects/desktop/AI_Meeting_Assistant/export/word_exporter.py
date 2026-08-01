from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


class WordExporter:

    def export(
            self,
            output_file: str,
            meeting_title: str,
            summary: str,
            transcript: str,
            actions: list = None,
            decisions: list = None):

        doc = Document()

        # 標題
        title = doc.add_heading(
            meeting_title,
            level=1
        )

        title.alignment = (
            WD_PARAGRAPH_ALIGNMENT.CENTER
        )

        # 會議摘要
        doc.add_heading(
            "會議摘要",
            level=2
        )

        p = doc.add_paragraph(
            summary
        )

        p.style.font.size = Pt(11)

        # 決策事項
        if decisions:

            doc.add_heading(
                "決策事項",
                level=2
            )

            for item in decisions:

                decision = item.get(
                    "decision",
                    ""
                )

                doc.add_paragraph(
                    f"• {decision}"
                )

        # 待辦事項
        if actions:

            doc.add_heading(
                "待辦事項",
                level=2
            )

            for item in actions:

                owner = item.get(
                    "owner",
                    ""
                )

                task = item.get(
                    "task",
                    ""
                )

                deadline = item.get(
                    "deadline",
                    ""
                )

                doc.add_paragraph(
                    f"• {owner}：{task} "
                    f"(截止：{deadline})"
                )

        # 逐字稿
        doc.add_heading(
            "逐字稿",
            level=2
        )

        doc.add_paragraph(
            transcript
        )

        Path(output_file).parent.mkdir(
            exist_ok=True
        )

        doc.save(
            output_file
        )

        return output_file


_exporter = WordExporter()


def export_word(
        output_file,
        meeting_title,
        summary,
        transcript,
        actions=None,
        decisions=None):

    return _exporter.export(
        output_file,
        meeting_title,
        summary,
        transcript,
        actions,
        decisions
    )